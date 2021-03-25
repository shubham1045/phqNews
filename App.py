from docx import Document
from docx.shared import Inches
import os
import requests
from bs4 import BeautifulSoup
import urllib.request
from tkinter import *
import urllib.request
import random
from docx import Document
#from docx.shared import Ibnches

import os
import requests
from bs4 import BeautifulSoup
import urllib.request
from tkinter import *
import urllib.request
import random
from docx import Document
from docx.shared import Inches
import datetime
from datetime import date
import time



def callback():
    print("click!")



var2=str("+")
today=date.today()
today=date.today()
filetime=str(datetime.datetime.now())

cnt=0



var2=str("+")





#---------------------CLASS  DOC ----------------------------------
class Doc:
    var = str(" ")
    document = Document()
    document.add_heading('NEWS REPORT: ' + var, 0)
    link1 = str(" ")

    def __init__(self, var1):
        Doc.var = var1
#   ==============F I R S T P O S T . C O M============================
    def fpa(self, url):


        resp = requests.get(url)
        time.sleep(5)
        if resp.status_code == 200:
            print("Successfully opened the web page")
            print("The news are as follow :-\n")

            soup = BeautifulSoup(resp.text, 'html.parser')

            l = soup.find("div", {"class": "article-full-content"})
            for i in l.findAll("p"):
                print(i.text)
                Doc.document.add_paragraph(
                    i.text
                )
                txt = str(i.text)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)

        else:
            print("Error")
    def fpl(self):
        var = str(var1.get())
        var3 = str(var2.get())
        # var = str(var1.get())
        resp = urllib.request.urlopen("https://www.firstpost.com/search/"+var +"%20"+var3)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            # if ((link['href'].find('altnews') == -1) and
            if ((len(link['href']) > 80)):
                # if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (link['href'].find('youtube')== -1) and (link['href'].find("ram") != -1)):
                print(link['href'])
                p = Doc.document.add_paragraph( style='List Number'

                                               )
                p.add_run(link['href']).bold = True
                link1 = (link['href'])
                txt=str(link1)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)
                try:
                    self.fpa(link1)
                except:
                    print("Article Removed by Author")

                # self.news(link['href'])
            #Doc.document.add_page_break()

            Doc.document.save(str(today)+str(var1.get()+var2.get())+".docx")

    #========================= A L T N E W S . C O M=================
    def alta(self, url):

            resp = requests.get(url)
            time.sleep(5)
            if resp.status_code == 200:
                print("Successfully opened the web page")
                print("The news are as follow :-\n")

                soup = BeautifulSoup(resp.text, 'html.parser')

                l = soup.find("div", {"class": "entry-content herald-entry-content"})
                for i in l.findAll("p"):
                    print(i.text)
                    Doc.document.add_paragraph(
                        i.text
                    )
                    txt = str(i.text)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)

            else:
                print("Error")

        # ============= ALT LINKS===============
    def altl(self):
            var = str(var1.get())
            var3 = str(var2.get())
            resp = urllib.request.urlopen(
                "https://news.google.com/search?q=" + var + "%20" + var3 + "%20when%3A1d%20site%3Aaltnews.in&hl=en-IN&gl=IN&ceid=IN%3Aen")
            soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
            for link in soup.find_all('a', href=True):
                # if ((link['href'].find('altnews') == -1) and
                if ((len(link['href']) > 80)):
                    if ((len(link['href']) > 50) and (link['href'].find('doc') == -1) and (
                            link['href'].find('article') != -1) and (link['href'].find('facebook') == -1) and (
                            link['href'].find('linkedin') == -1) and (link['href'].find('youtube') == -1) and (
                            link['href'].find('twitter') == -1) and (link['href'].find('play') == -1) and (
                            link['href'].find('plus') == -1) and (link['href'].find('google') == -1) and (
                            link['href'].find('topics') == -1) and (link['href']).find('translate') == -1 and (
                    link['href']).find('recap') == -1 and (link['href']).find('highlights') == -1):

                        print(link['href'])
                        p = Doc.document.add_paragraph(style='List Number'
                                                       )
                        p.add_run(link['href']).bold = True
                        link1 = (link['href'])
                        # txt = str(link1)
                        txt = str(link1)
                        T = Text(top, height=5, width=60)
                        T.place(x=20, y=100)
                        T.insert(END, txt)
                        try:
                            self.alta(link1)
                        except:
                            print("Content Removed")
                Doc.document.save(str(today)+str(var1.get()+var2.get())+".docx")

    #========= A N I . C O M=================================
    def anil(self):
            var = str(var1.get())
            var3= str(var2.get())
            resp = urllib.request.urlopen("https://www.aninews.in/search/?query=" + var+"-"+var3+"/")
            soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
            for link in soup.find_all('a', href=True):
                if ((link['href'].find('http') == -1) and (len(link['href']) > 50) and (link['href'].find(var) != -1)):
                    if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1)):
                        print("https://www.aninews.in"+link['href'])
                        p=Doc.document.add_paragraph( style='List Number'
                        )
                        p.add_run(("https://www.aninews.in"+
                            link['href'])).bold = True
                        link1 = (link['href'])
                        #self.dba(link1)
                        txt = str(link1)
                        T = Text(top, height=5, width=60)
                        T.place(x=20, y=100)
                        T.insert(END, txt)
                    # self.news(link['href'])
                    # Doc.document.add_page_break()

                    Doc.document.save(str(today)+str(var1.get()+var2.get())+".docx")
#============= N A V H  I N D T I M E S . I N    LINKS===========
    def navhinda(self, url):

        resp = requests.get(url)
        time.sleep(5)
        if resp.status_code == 200:
            print("Successfully opened the web page")
            print("The news are as follow :-\n")

            soup = BeautifulSoup(resp.text, 'html.parser')

            l = soup.find("div", {"class": "entry"})
            for i in l.findAll("p"):
                print(i.text)
                Doc.document.add_paragraph(
                    i.text
                )
                txt = str(i.text)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)

        else:
            print("Error")
        # ============= Navhind times .com =========

    def navhindl(self):
        var = str(var1.get())
        var3 = str(var2.get())
        resp = urllib.request.urlopen(
            "https://news.google.com/search?q=" + var + "%20" + var3 + "%20when%3A1d%20site%3Anavhindtimes.in&hl=en-IN&gl=IN&ceid=IN%3Aen")
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            # if ((link['href'].find('altnews') == -1) and
            if ((len(link['href']) > 80)):
                if ((len(link['href']) > 50) and (link['href'].find('doc') == -1) and (
                        link['href'].find('article') != -1) and (link['href'].find('facebook') == -1) and (
                        link['href'].find('linkedin') == -1) and (link['href'].find('youtube') == -1) and (
                        link['href'].find('twitter') == -1) and (link['href'].find('play') == -1) and (
                        link['href'].find('plus') == -1) and (link['href'].find('google') == -1) and (
                        link['href'].find('topics') == -1) and (link['href']).find('translate') == -1 and (
                link['href']).find('recap') == -1 and (link['href']).find('highlights') == -1):

                    print(link['href'])
                    url = str("https://news.google.com" + link['href'])
                    p = Doc.document.add_paragraph(style='List Number'
                                                   )
                    p.add_run(url).bold = True
                    link1 = (link['href'])
                    # txt = str(link1)

                    txt = str(url)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    try:
                        self.navhinda(url)
                    except:
                        print("Content Removed")
            Doc.document.save(str(today)+str(var1.get()+var2.get())+".docx")
#============ SIFY .COM LINKS================
    def sifyl(self):
        var = str(var1.get())
        var3 = str(var2.get())
        # var = str(var1.get())
        sify=0;
        resp = urllib.request.urlopen("http://www.sify.com/")
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            # if ((link['href'].find('altnews') == -1) and
            if ((len(link['href']) > 50)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (
                        link['href'].find('youtube') == -1) and (link['href'].find("sify") != -1)):
                    print(link['href'])
                    p=Doc.document.add_paragraph( style='List Number'
                                               )
                    p.add_run(link['href']).bold=True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    try:
                        if(sify<5):
                            self.sify_article(link1)
                            sify+=1
                    except:
                        print("Article removed for the website:: ")
            Doc.document.save(str(today) + str(var1.get() + var2.get()) + ".docx")

    #================= SIFY . COM   ARTICLE ===============
    def sify_article(self,url):
        #url = 'https://www.sify.com/news/army-says-youth-killed-in-pulwama-not-a-soldier-news-national-tdnrEsdefcfcj.html'

        resp = requests.get(url)

        if resp.status_code == 200:
            print("Successfully opened the web page")
            print("The news are as follow :-\n")

            soup = BeautifulSoup(resp.text, 'html.parser')

            l = soup.find("div", {"class": "fullstory-txt"})
            for i in l.findAll("p"):
                print(i.text)
                docx.document.add_paragraph(i.text)
                txt = str(i.text)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)
        else:
            print("Error")
#================= N D T V . C O M ==============
    def ndtvl(self):
        var3 = str(var2.get())
        var = str(var1.get())
        ndtv=0
        resp = urllib.request.urlopen("https://www.ndtv.com/topic/"+var+"-"+var3)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            # if ((link['href'].find('altnews') == -1) and
            if ((len(link['href']) > 80)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (
                        link['href'].find('facebook') == -1) and (link['href'].find('linkedin') == -1) and (
                        link['href'].find('youtube') == -1) and (link['href'].find('plus') == -1) and (
                        link['href'].find("ndtv") != -1)):
                    print(link['href'])
                    p=Doc.document.add_paragraph( style='List Number'
                                               )
                    p.add_run(link['href']).bold=True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    try:
                        if(ndtv<5):
                            self.ndtv_article(link1)
                            ndtv+=1
                    except:
                        print("Article removed for the website:: ")
                Doc.document.save(str(today)+str(var1.get()+var2.get())+".docx")
#888888888888888888
    def indian_expressl(self):
        var = str(var1.get())
        var3= str(var2.get())
        ie=0
        resp = urllib.request.urlopen("https://indianexpress.com/?s="+var+"+"+var3)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            # if ((link['href'].find('altnews') == -1) and
            if ((len(link['href']) > 50)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (
                        link['href'].find('facebook') == -1) and (link['href'].find('linkedin') == -1) and (
                        link['href'].find('youtube') == -1) and (link['href'].find('plus') == -1) and (
                        link['href'].find("indian") != -1)):
                    print(link['href'])
                    p=Doc.document.add_paragraph( style='List Number'
                                               )
                    p.add_run(link['href']).bold = True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    try:

                        if(ie<5):
                            self.indian_express_article(link1)
                            ie+=1
                    except:
                        print("Article removed for the website:: ")
            Doc.document.save(str(today) + str(var1.get() + var2.get()) + ".docx")
#=================== I N D I A N    E X P R E S S ===============
    def indian_expressll(self):
        var = str(var1.get())
        var3= str(var2.get())
        resp = urllib.request.urlopen("https://indianexpress.com/?s="+var+"+"+var3)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            # if ((link['href'].find('altnews') == -1) and
            if ((len(link['href']) > 50)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (
                        link['href'].find('facebook') == -1) and (link['href'].find('linkedin') == -1) and (
                        link['href'].find('youtube') == -1) and (link['href'].find('plus') == -1) and (
                        link['href'].find("indian") != -1)):
                    print(link['href'])
                    p=Doc.document.add_paragraph( style='List Number'
                                               )
                    p.add_run(link['href']).bold = True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    try:
                        self.indian_express_article(link1)
                    except:
                        print("Article removed for the website:: ")
                Doc.document.save(str(var1.get() + var2.get()) + ".docx")
        def opendocx():
            os.startfile("G:/Projects Python/News/" + var + var3 + ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():

            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var) +                           str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()
#============ INDIAN EXPRESS . COM  ARTICLE=============
    def indian_express_article(self,url):
        #url = 'https://indianexpress.com/article/opinion/columns/congress-secularism-lok-sabha-elections-2019-5623316/'

        resp = requests.get(url)

        if resp.status_code == 200:
            print("Successfully opened the web page")
            print("The news are as follow :-\n")

            soup = BeautifulSoup(resp.text, 'html.parser')

            l = soup.find("div", {"class": "o-story-content__main a-wysiwyg"})
            for i in l.findAll("p"):
                print(i.text)
                docx.document.add_paragraph(i.text)
                txt = str(i.text)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)
        else:
            print("Error")
#================= N D T V . C O M  ARTICLE===============
    def ndtv_article(self,url):
        #url = 'https://www.ndtv.com/india-news/ayodhya-case-will-be-referred-for-mediation-says-supreme-court-2004516'

        resp = requests.get(url)

        if resp.status_code == 200:
            print("Successfully opened the web page")
            print("The news are as follow :-\n")

            soup = BeautifulSoup(resp.text, 'html.parser')

            l = soup.find("div", {"class": "ins_storybody"})
            for i in l.findAll("p"):
                print(i.text)
                docx.document.add_paragraph(i.text)
                txt = str(i.text)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)
        else:
            print("Error")
    #================== S C R O L L . I N ================
    def scll(self):
        var3=str(var2.get())
        var = str(var1.get())
        scl=0
        resp = urllib.request.urlopen("https://scroll.in/search?q="+var+"%20"+var3+"&page=1")
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            # if ((link['href'].find('altnews') == -1) and
            if ((len(link['href']) > 80)):
                # if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (link['href'].find('youtube')== -1) and (link['href'].find("ram") != -1)):
                print(link['href'])
                p=Doc.document.add_paragraph( style='List Number'
                                           )
                p.add_run(link['href']).bold= True
                link1 = (link['href'])
                txt = str(link1)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)
                try:
                    if(scl<5):
                        self.scla(link1)
                        scl+=1
                except:
                    print("Article removed for the website:: ")
            Doc.document.save(str(today)+str(var1.get()+var2.get())+".docx")
#-----------d
    def sclll(self):
        var3=str(var2.get())
        var = str(var1.get())
        resp = urllib.request.urlopen("https://scroll.in/search?q="+var+"%20"+var3+"&page=1")
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            # if ((link['href'].find('altnews') == -1) and
            if ((len(link['href']) > 80)):
                # if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (link['href'].find('youtube')== -1) and (link['href'].find("ram") != -1)):
                print(link['href'])
                p=Doc.document.add_paragraph( style='List Number'
                                           )
                p.add_run(link['href']).bold= True
                link1 = (link['href'])
                txt = str(link1)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)
                try:
                    self.scla(link1)
                except:
                    print("Article removed for the website:: ")
            Doc.document.save(str(var1.get() + var2.get()) + ".docx")

        def opendocx():
            os.startfile("G:/Projects Python/News/" + var + var3 + ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var) + str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()


#================== S C R O L L . I N   ARTICLE ====================
    def scla(self,url):
        #url = 'https://scroll.in/article/915890/what-the-manifestations-of-navaratri-tell-us-about-religious-expression-and-narratives-of-power'

        resp = requests.get(url)

        if resp.status_code == 200:
            print("Successfully opened the web page")
            print("The news are as follow :-\n")

            soup = BeautifulSoup(resp.text, 'html.parser')

            l = soup.find("div", {"class": "article-body"})
            for i in l.findAll("p"):
                print(i.text)
                docx.document.add_paragraph(i.text
                                       )
                txt = str(i.text)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)
                #Doc.document.add_page_break()
        else:
            print("Error")

        # ================DAINIK BHASKAR LINK FETCHER==================
    def dbl(self):
            var = str(var1.get())
            var3= str(var2.get())
            db=0
            resp = urllib.request.urlopen("https://www.bhaskar.com/topics/" + var+"-"+var3+"/")
            soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
            for link in soup.find_all('a', href=True):
                if ((link['href'].find('http') != -1) and (len(link['href']) > 50) and (link['href'].find(var) != -1)):
                    if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1)):
                        print(link['href'])
                        Doc.document.add_paragraph(
                            link['href'], style='List Number'
                        )
                        link1 = (link['href'])
                        txt = str(link1)
                        T = Text(top, height=5, width=60)
                        T.place(x=20, y=100)
                        T.insert(END, txt)
                        try:
                            if(db<5):
                                self.dba(link1)
                                db+=1
                        except:
                            print("No permission from the author")

                    # self.news(link['href'])
                    # Doc.document.add_page_break()

                Doc.document.save(str(today) + str(var1.get() + var2.get()) + ".docx")
    #=============DAINIK BHASKAR ARTICLE FETCHER====================
    def dba(self, url):


        resp = requests.get(url)

        if resp.status_code == 200:
            print("Successfully opened the web page")
            print("The news are as follow :-\n")

            soup = BeautifulSoup(resp.text, 'html.parser')

            l = soup.find("div", {"class": "db_storycontent"})
            for i in l.findAll("p"):
                print(i.text)
                Doc.document.add_paragraph(
                    i.text
                )
                txt = str(i.text)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)

        else:
            print("Error")

#============ F I R S T  P O S T . C O M============================





#=====================CENTRAL CHRONICAL ARTICLE FINDER==============
    def cca(self, url):

        resp = requests.get(url)

        if resp.status_code == 200:
            print("Successfully opened the web page")
            print("The news are as follow :-\n")

            soup = BeautifulSoup(resp.text, 'html.parser')

            l = soup.find("div", {"class": "entry-content"})
            for i in l.findAll("p"):
                print(i.text)
                Doc.document.add_paragraph(
                    i.text
                )
                txt = str(i.text)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)


        else:
            print("Error")

        # ----------------------------CentralCHronical IN LINKS----------------------

    def ccl(self):
        var = str(var1.get())
        var3 = str(var2.get())
        cc=0
        resp = urllib.request.urlopen("http://www.centralchronicle.com/?s="+var+"+"+var3)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 50)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (
                        link['href'].find('youtube') == -1) and (link['href'].find(str(var)) != -1)):
                    print(link['href'])
                    p = Doc.document.add_paragraph(style='List Number'

                                                   )
                    p.add_run(link['href']).bold = True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    try:
                        if(cc<5):
                            self.cca(link1)
                            cc+=1
                    except:
                        print("Permission Denied")

                # self.news(link['href'])
                #Doc.document.add_page_break()

            Doc.document.save(str(today) + str(var1.get() + var2.get()) + ".docx")
#============== Chrnical for button
    def ccll(self):
        var = str(var1.get())
        var3 = str(var2.get())
        resp = urllib.request.urlopen("http://www.centralchronicle.com/?s="+var+"+"+var3)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 50)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (
                        link['href'].find('youtube') == -1) and (link['href'].find(str(var)) != -1)):
                    print(link['href'])
                    p = Doc.document.add_paragraph(style='List Number'

                                                   )
                    p.add_run(link['href']).bold = True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    try:
                        self.cca(link1)
                    except:
                        print("Permission Denied")

                # self.news(link['href'])
                #Doc.document.add_page_break()

                Doc.document.save(str(var1.get()+var2.get()) + ".docx")
        def opendocx():
            os.startfile("G:/Projects Python/News/" + var + var3 + ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var) +                          str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()
    #==========================CAMERA24.IN ARTICLE FETCHER==================
    def c4a(self, url):

        resp = requests.get(url)

        if resp.status_code == 200:
            print("Successfully opened the web page")
            print("The news are as follow :-\n")

            soup = BeautifulSoup(resp.text, 'html.parser')

            l = soup.find("div", {"class": "entry-content"})
            for i in l.findAll("p"):
                print(i.text)
                Doc.document.add_paragraph(
                    i.text
                )
                txt = str(i.text)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)


        else:
            print("Error")

        # ----------------------------CAMERA24.IN LINKS----------------------

    def c4l(self):
        var = str(var1.get())
        resp = urllib.request.urlopen("https://camera24.in/")
        c4=1;
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 50) and (link['href'].find(var) != -1)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1)):
                    print(link['href'])
                    p = Doc.document.add_paragraph(link['href'],style='List Number'

                                                   )
                    p.add_run(link['href']).bold = True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    if(c4<3):
                        self.c4a(link1)
                        c4+=1


                # self.news(link['href'])
                Doc.document.add_page_break()

                Doc.document.save(str(var1.get()+var2.get()) + ".docx")
#----------- C24 links for button -------
    def c4ll(self):
        var = str(var1.get())
        var3 = str(var2.get())
        resp = urllib.request.urlopen("https://camera24.in/")
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 50) and (link['href'].find(var) != -1)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1)):
                    print(link['href'])
                    p = Doc.document.add_paragraph(link['href'],style='List Number'

                                                   )
                    p.add_run(link['href']).bold = True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    self.c4a(link1)


                # self.news(link['href'])
                Doc.document.add_page_break()

                Doc.document.save(str(var1.get()+var2.get()) + ".docx")
        def opendocx():
            os.startfile("G:/Projects Python/News/" + var + var3 + ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var) +                          str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()
#=======================AAJTAK ARTICLE FETCHER==================
    def aaja(self, url):

        resp = requests.get(url)

        if resp.status_code == 200:
            print("Successfully opened the web page")
            print("The news are as follow :-\n")

            soup = BeautifulSoup(resp.text, 'html.parser')

            l = soup.find("div", {"class": "detailTxtContainer storyBody middle_s adblockcontainer"})
            for i in l.findAll("p"):
                print(i.text)
                Doc.document.add_paragraph(
                    i.text
                )
                txt = str(i.text)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)


        else:
            print("Error")

# ----------------------------AAJTAK LINKS----------------------

    def aajl(self):
        var = str(var1.get())
        var3 = str(var2.get())
        resp = urllib.request.urlopen("https://aajtak.intoday.in/topic/"+var+"-"+var3+"-news.html?datestart="+str(today.day-1)+"/04/2019&dateend="+str(today.day)+"/04/2019")
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 80) and (link['href'].find(var) != -1)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1)and (link['href'].find('aajtak') != -1)):
                    print(link['href'])
                    p = Doc.document.add_paragraph(style='List Number'

                                                   )
                    p.add_run(link['href']).bold = True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    try:
                        self.aaja(link1)
                    except:
                        print("Permission denied from aajtak")

                # self.news(link['href'])
                #Doc.document.add_page_break()

            Doc.document.save(str(today) + str(var1.get() + var2.get()) + ".docx")
#============== AAJ TAK LINKS FOR BUTTON =============
    def aajll(self):
        var = str(var1.get())
        var3 = str(var2.get())
        resp = urllib.request.urlopen("https://aajtak.intoday.in/topic/"+var+"-"+var3+".html")
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 80) and (link['href'].find(var) != -1)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1)and (link['href'].find('aajtak') != -1)):
                    print(link['href'])
                    p = Doc.document.add_paragraph(style='List Number'

                                                   )
                    p.add_run(link['href']).bold = True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    try:
                        self.aaja(link1)
                    except:
                        print("Permission denied from aajtak")

                # self.news(link['href'])
                #Doc.document.add_page_break()

                Doc.document.save(str(var1.get()+var2.get()) + ".docx")
        def opendocx():
            os.startfile("G:/Projects Python/News/" + var + var3 + ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var) +                           str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()

    #===============PATRIKA NEWS FETCHER===================
    def pka(self, url):

        resp = requests.get(url)

        if resp.status_code == 200:
            print("Successfully opened the web page")
            print("The news are as follow :-\n")

            soup = BeautifulSoup(resp.text, 'html.parser')

            l = soup.find("div", {"class": "complete-story"})
            for i in l.findAll("p"):
                print(i.text)
                Doc.document.add_paragraph(
                    i.text
                )
                txt = str(i.text)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)


        else:
            print("Error")
#===================SUDARSHAN NEWS ARTICLE ================
    def ssna(self, url):

        resp = requests.get(url)

        if resp.status_code == 200:
            print("Successfully opened the web page")
            print("The news are as follow :-\n")

            soup = BeautifulSoup(resp.text, 'html.parser')

            l = soup.find("div", {"class": "entry-content clearfix"})
            for i in l.findAll("p"):
                print(i.text)
                Doc.document.add_paragraph(
                    i.text
                )
                txt = str(i.text)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)
                Doc.document.save(str(var1.get() + var2.get()) + ".docx")


        else:
            print("Error")
#==================SUDARSHAN NEWS . COM-----------------------------
    def ssnl(self):
        var = str(var1.get())
        var3 = str(var2.get())
        ss=0
        resp = urllib.request.urlopen("http://www.sudarshannews.in/?s=" + var + "+" + var3)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 50) and (link['href'].find(var) != -1)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (
                        link['href'].find('youtube') == -1) and (link['href'].find('sudarshannews') != -1) and (
                        link['href'].find('addtoany') == -1)):
                    print(link['href'])
                    p = Doc.document.add_paragraph(style='List Number'

                                               )
                    p.add_run(link['href']).bold = True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    try:
                        if(ss<3):
                            self.ssna(link1)
                            ss+=1;
                    except:
                        print("System detects you as a robot ")

                # self.news(link['href'])
                # Doc.document.add_page_break()

            Doc.document.save(str(var1.get()+var2.get()) + ".docx")
#777777777777777777
    def ssnll(self):
        var = str(var1.get())
        var3 = str(var2.get())
        resp = urllib.request.urlopen("http://www.sudarshannews.in/?s=" + var + "+" + var3)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 50) and (link['href'].find(var) != -1)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (
                        link['href'].find('youtube') == -1) and (link['href'].find('sudarshannews') != -1) and (
                        link['href'].find('addtoany') == -1)):
                    print(link['href'])
                    p = Doc.document.add_paragraph(style='List Number'

                                               )
                    p.add_run(link['href']).bold = True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    try:
                        self.ssna(link1)
                    except:
                        print("System detects you as a robot ")

                # self.news(link['href'])
                # Doc.document.add_page_break()

            Doc.document.save(str(var1.get()+var2.get()) + ".docx")

        def opendocx():
            os.startfile("G:/Projects Python/News/" + var + var3 + ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var) + str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()

        # ----------------------------PATRIKA NEWS + LINKS----------------------

    def pkl(self):
        var = str(var1.get())
        var3 = str(var2.get())
        pk=0
        resp = urllib.request.urlopen("https://www.patrika.com/search/?q=" + var+"+"+var3)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 50) and (link['href'].find(var) != -1)):
                #if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1)):
                    print(link['href'])
                    p = Doc.document.add_paragraph(style='List Number'

                                                   )
                    p.add_run(link['href']).bold = True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)

                    try:
                         if(pk<4):
                             self.pka(link1)
                             pk+=1
                    except:
                         print("No access")

                # self.news(link['href'])
                #Doc.document.add_page_break()

            Doc.document.save(str(today)+str(var1.get()+var2.get())+".docx")
#--------------PKLLinks----------
    def pkll(self):
        var = str(var1.get())
        var3 = str(var2.get())
        resp = urllib.request.urlopen("https://www.patrika.com/search/?q=" + var+"+"+var3)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 50) and (link['href'].find(var) != -1)):
                #if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1)):
                    print(link['href'])
                    p = Doc.document.add_paragraph(style='List Number'

                                                   )
                    p.add_run(link['href']).bold = True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)

                    try:
                         self.pka(link1)
                    except:
                         print("No access")

                # self.news(link['href'])
                #Doc.document.add_page_break()
            Doc.document.save(str(var1.get()+var2.get()) + ".docx")
        def opendocx():
            os.startfile("G:/Projects Python/News/" + var + var3 + ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var) + str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()
  #===================The HINDU
    def thl(self):
        var = str(var1.get())
        var3 = str(var2.get())
        resp = urllib.request.urlopen("https://www.thehindu.com/search/?q="+var+"+"+var3+"&order=DESC&sort=publishdate")
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 50) and (link['href'].find(var)!=-1) ):
                if((link['href'].find('whatsapp') == -1) and (link['href'].find('video')== -1)):
                    print(link['href'])
                    Doc.document.add_paragraph(
                        link['href'], style='List Number'
                    )
            Doc.document.save(str(today)+str(var1.get()+var2.get())+".docx")
            link1 = (link['href'])
            txt = str(link1)
            T = Text(top, height=5, width=60)
            T.place(x=20, y=100)
            T.insert(END, txt)
                    #self.dba(link1)

#==============HINDUSTAN TIMES ARTICLE
    def hta(self, url):


        resp = requests.get(url)

        if resp.status_code == 200:
            print("Successfully opened the web page")
            print("The news are as follow :-\n")

            soup = BeautifulSoup(resp.text, 'html.parser')

            l = soup.find("div", {"class": "story-details"})
            for i in l.findAll("p"):
                print(i.text)
                Doc.document.add_paragraph(
                    i.text
                )
                txt = str(i.text)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)


        else:
            print("Error")

 #==================HINDUSTAN TIMES LINKFPA
    def htl(self):
        # var = str(var1.get())
        var = str(var1.get())
        var3 = str(var2.get())
        url = "https://www.bing.com/news/search?q=site%3ahindustantimes.com+" + var + var3 + "&qft=interval%3d%227%22&form=PTFTNR"
        resp = urllib.request.urlopen(url)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            # if ((link['href'].find('altnews') == -1) and
            if ((len(link['href']) > 80) and (link['href'].find('hindustantimes') != -1) and (
                    link['href'].find('whatsapp') == -1) and (link['href'].find('facebook') == -1) and (
                    link['href'].find('linkedin') == -1) and (link['href'].find('youtube') == -1) and (
                    link['href'].find('twitter') == -1) and (link['href'].find('play') == -1) and (
                    link['href'].find('plus') == -1) and (link['href'].find('www') != -1) and (
                    link['href'].find('microsoft') == -1) and (link['href']).find('pragya') != -1):
                print(link['href'])
                Doc.document.add_paragraph(
                    link['href'], style='List Number'
                )
                link1 = (link['href'])
                txt = str(link1)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)
                try:
                    self.hta(link1)
                except:
                    print("No permission from the author")

                # self.news(link['href'])
                # Doc.document.add_page_break()

            Doc.document.save(str(today) + str(var1.get() + var2.get()) + ".docx")

#========= Hindustan Times Link ==============
    def htll(self):
        var = str(var1.get())
        var3 = str(var2.get())
        resp = urllib.request.urlopen("https://www.bing.com/news/search?q=site%3Ahindustantimes.com++" + var+"+"+var3)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 50) and (link['href'].find(var)!=-1) ):
                if((link['href'].find('whatsapp') == -1) and (link['href'].find('video')== -1)and link['href'].find('twitter')==-1):
                    print(link['href'])
                    p=Doc.document.add_paragraph( link['href'],style='List Number'

                    )
                    p.add_run(link['href']).bold = True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    try:
                        self.hta(link1)
                    except:
                        print("Article fetching is restricted")


                #self.news(link['href'])
                Doc.document.add_page_break()

                Doc.document.save(str(var1.get()+var2.get())+".docx")
            Doc.document.save(str(var1.get() + var2.get()) + ".docx")

        def opendocx():
            os.startfile("G:/Projects Python/News/" + var + var3 + ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var) + str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()

        popupmsg(var)
#========== HINDI FIRST POST ARTICLE++++++++++++++++++
    def hfpa(self, url):


        resp = requests.get(url)

        if resp.status_code == 200:
            print("Successfully opened the web page")
            print("The news are as follow :-\n")

            soup = BeautifulSoup(resp.text, 'html.parser')

            l = soup.find("div", {"class": "csmpn clearfix"})
            for i in l.findAll("p"):
                print(i.text)
                Doc.document.add_paragraph(
                    i.text
                )
                txt = str(i.text)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)


        else:
            print("Error")
#================= HINDI FIRST POST WITH SEARCH============
    def hfpl(self):
        var = str(var1.get())
        var3 = str(var2.get())
        hf=0
        resp = urllib.request.urlopen("https://hindi.firstpost.com/search/"+var+"%20"+var3+"/")
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((len(link['href']) > 50)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (
                        link['href'].find('facebook') == -1) and (link['href'].find('linkedin') == -1) and (
                        link['href'].find('youtube') == -1) and (link['href'].find('twitter') == -1) and (
                        link['href'].find('plus') == -1) and (link['href'].find(var) != -1)):
                    print(link['href'])
                    p=Doc.document.add_paragraph( link['href'],style='List Number'

                    )
                    p.add_run(link['href']).bold = True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    try:
                        if(hf<4):
                            self.hfpa(link1)
                            hf+=1
                    except:
                        print("Article fetching is restricted")


                #self.news(link['href'])
                #Doc.document.add_page_break()

                Doc.document.save(str(var1.get()+var2.get())+".docx")


    #=============MAIN FUNCTION FETCH LINK===================
    def fetchlink(self):
        var = str(var1.get())
        var3 = str(var2.get())
        resp = urllib.request.urlopen("https://www.bing.com/news/search?q=" + var+"+"+var3 + "&FORM=NWRFSH")
        # resp = urllib.request.urlopen("https://www.google.co.in/search?q="+var+"&source=lnms&tbm=nws")
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))

        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 50) and (link['href'].find(var)!=1)):
                print(link['href'])
                txt=link['href']
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)
                Doc.document.add_paragraph(
                    link['href'], style='List Number'
                )
                Doc.document.save(str(var1.get() + var2.get()) + ".docx")
        def opendocx():
            os.startfile("G:/Projects Python/News/" + var+var3+ ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var)+str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()

        popupmsg(var)



#============= B B C LINK FETCHER ============
    def bbcl(self):
        var = str(var1.get())
        var3 = str(var2.get())
        resp = urllib.request.urlopen("https://www.bing.com/news/search?q=site%3Abbc.com+"+var+"&q")
        # resp = urllib.request.urlopen("https://www.google.co.in/search?q="+var+"&source=lnms&tbm=nws")
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))

        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 50) and (link['href'].find(var)!=1)):
                print(link['href'])
                Doc.document.add_paragraph(
                    link['href'], style='List Number'
                )
                Doc.document.save(str(var1.get() + var2.get()) + ".docx")
        def opendocx():
            os.startfile("G:/Projects Python/News/" + var+var3+ ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var)+str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()

        popupmsg(var)
#=========== NDTV AA ============
    def ndtva_article(self,url):
        #url = 'https://www.ndtv.com/india-news/ayodhya-case-will-be-referred-for-mediation-says-supreme-court-2004516'

        resp = requests.get(url)

        if resp.status_code == 200:
            print("Successfully opened the web page")
            print("The news are as follow :-\n")

            soup = BeautifulSoup(resp.text, 'html.parser')

            l = soup.find("div", {"class": "ins_storybody"})
            for i in l.findAll("p"):
                print(i.text)
                docx.document.add_paragraph(i.text)
                txt = str(i.text)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)
        else:
            print("Error")
#===============NDTV lllLLL=========
    def ndtvll(self):
        var3 = str(var2.get())
        var = str(var1.get())
        resp = urllib.request.urlopen("https://www.ndtv.com/topic/" + var + "-" + var3)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
                # if ((link['href'].find('altnews') == -1) and
            if ((len(link['href']) > 80)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (
                        link['href'].find('facebook') == -1) and (link['href'].find('linkedin') == -1) and (
                        link['href'].find('youtube') == -1) and (link['href'].find('plus') == -1) and (
                        link['href'].find("ndtv") != -1)):

                    print(link['href'])
                    p = Doc.document.add_paragraph(style='List Number'
                                                       )
                    p.add_run(link['href']).bold = True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    try:
                        self.ndtv_article(link1)
                    except:
                        print("Article removed for the website:: ")
                Doc.document.save(str(var1.get() + var2.get()) + ".docx")
        def opendocx():
            os.startfile("G:/Projects Python/News/" + var+var3+ ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var)+str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()

        popupmsg(var)
#==========FIRST POST LINK ===========
    def fpll(self):
        var = str(var1.get())
        var3 = str(var2.get())
        # var = str(var1.get())
        resp = urllib.request.urlopen("https://www.firstpost.com/search/" + var + "%20" + var3)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
                # if ((link['href'].find('altnews') == -1) and
            if ((len(link['href']) > 80)):
                    # if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (link['href'].find('youtube')== -1) and (link['href'].find("ram") != -1)):
                print(link['href'])
                p = Doc.document.add_paragraph(style='List Number'

                                                   )
                p.add_run(link['href']).bold = True
                link1 = (link['href'])
                txt = str(link1)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)
                try:
                    self.fpa(link1)
                except:
                    print("Article Removed by Author")

                    # self.news(link['href'])
                # Doc.document.add_page_break()
            self.hfpl
            Doc.document.save(str(var1.get() + var2.get()) + ".docx")

        def opendocx():
            os.startfile("G:/Projects Python/News/" + var+var3+ ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var)+str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()

#   =   =   =   NavHind Times Links fro Button =    =   =   =   =
    def navhindll(self):
        var = str(var1.get())
        var3 = str(var2.get())
        # var = str(var1.get())
        resp = urllib.request.urlopen("http://www.navhindtimes.in/?s="+var+"+"+var3)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            # if ((link['href'].find('altnews') == -1) and
            if ((len(link['href']) > 50)):
                # if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (link['href'].find('youtube')== -1) and (link['href'].find("ram") != -1)):
                print(link['href'])

                p=Doc.document.add_paragraph( style='List Number'
                                           )
                p.add_run(link['href']).bold = True
                link1 = (link['href'])
                txt = str(link1)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)
                try:
                    self.navhindtimes_article(link1)
                except:
                    print("Articel removed for the website:: ")
            Doc.document.save(str(var1.get() + var2.get()) + ".docx")
        def opendocx():
            os.startfile("G:/Projects Python/News/" + var+var3+ ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var)+str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()
#========== SIFY LINKS FOR BUTTONS ===========
    def sifyll(self):
        var = str(var1.get())
        var3 = str(var2.get())
        # var = str(var1.get())
        resp = urllib.request.urlopen("http://www.sify.com/")
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            # if ((link['href'].find('altnews') == -1) and
            if ((len(link['href']) > 50)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (
                        link['href'].find('youtube') == -1) and (link['href'].find("sify") != -1)):
                    print(link['href'])
                    p=Doc.document.add_paragraph( style='List Number'
                                               )
                    p.add_run(link['href']).bold=True
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    try:
                        self.sify_article(link1)
                    except:
                        print("Article removed for the website:: ")
                Doc.document.save(str(var1.get() + var2.get()) + ".docx")

        def opendocx():
            os.startfile("G:/Projects Python/News/" + var + var3 + ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var) + str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()
#===============AlTL links===============
    def altll(self):
        var = str(var1.get())
        var3 = str(var2.get())
        resp = urllib.request.urlopen("https://www.altnews.in/?s=" + var + "+" + var3)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            # if ((link['href'].find('altnews') == -1) and
            if ((len(link['href']) > 80)):
                    # if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1) and (link['href'].find('youtube')== -1) and (link['href'].find("ram") != -1)):
                print(link['href'])
                p = Doc.document.add_paragraph(style='List Number'
                                               )
                p.add_run(link['href']).bold = True
                link1 = (link['href'])
                # txt = str(link1)
                txt = str(link1)
                T = Text(top, height=5, width=60)
                T.place(x=20, y=100)
                T.insert(END, txt)
                try:
                    self.dba(link1)
                except:
                    print("Content Removed")
            Doc.document.save(str(var1.get() + var2.get()) + ".docx")
        def opendocx():
            os.startfile("G:/Projects Python/News/" + var+var3+ ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var)+str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()

# ++++++++++++ DanikBhaskar Links for button +++++++++++
    def dbll(self):
        var = str(var1.get())
        var3 = str(var2.get())
        resp = urllib.request.urlopen("https://www.bhaskar.com/topics/" + var + "-" + var3 + "/")
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 50) and (link['href'].find(var) != -1)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1)):
                    print(link['href'])
                    Doc.document.add_paragraph(
                        link['href'], style='List Number'
                        )
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                    try:
                        self.dba(link1)
                    except:
                        print("No permission from the author")

                    # self.news(link['href'])
                    # Doc.document.add_page_break()

                Doc.document.save(str(var1.get() + var2.get()) + ".docx")

        def opendocx():
            os.startfile("G:/Projects Python/News/" + var + var3 + ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var) + str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()

#============= THE HINDU LINK EXTRACTER FOR BUTTONS ===========
    def thll(self):
        var = str(var1.get())
        var3 = str(var2.get())
        resp = urllib.request.urlopen(
                "https://www.thehindu.com/search/?q=" + var + "+" + var3 + "&order=DESC&sort=publishdate")
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 50) and (link['href'].find(var) != -1)):
                if ((link['href'].find('whatsapp') == -1) and (link['href'].find('video') == -1)):
                    print(link['href'])
                    Doc.document.add_paragraph(
                        link['href'], style='List Number'
                        )
                    Doc.document.save(str(var1.get() + var2.get()) + ".docx")
                    link1 = (link['href'])
                    txt = str(link1)
                    T = Text(top, height=5, width=60)
                    T.place(x=20, y=100)
                    T.insert(END, txt)
                Doc.document.save(str(var1.get() + var2.get()) + ".docx")

        def opendocx():
            os.startfile("G:/Projects Python/News/" + var + var3 + ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var) + str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()
        """
        Doc.document.add_page_break()
        docx.document.add_heading('NEWS REPORT: DainikBhaskar' + var, 0)
        self.dbl()
        Doc.document.add_page_break()
        docx.document.add_heading('NEWS REPORT: TheHindu' + var, 0)
        self.thl()
        Doc.document.add_page_break()
        docx.document.add_heading('NEWS REPORT: HindustanTimes' +var,  0)
        self.htl()
        Doc.document.add_page_break()
        docx.document.add_heading('NEWS REPORT: Camera24.in' +  var,0)
        self.c4l()
        Doc.document.add_page_break()
        docx.document.add_heading('NEWS REPORT: AajTak' + var,0)
        self.aajl()
        Doc.document.add_page_break()
        docx.document.add_heading('NEWS REPORT: firstpost' + var,0)
        self.fpl()
        Doc.document.add_page_break()
        docx.document.add_heading('NEWS REPORT: ALT NEWS' + var,0)
        self.altl()
        Doc.document.add_page_break()
        docx.document.add_heading('NEWS REPORT: ANI NEWS' +var, 0)
        self.anil()
        Doc.document.add_page_break()
        docx.document.add_heading('NEWS REPORT: NavHind' + var,0)
        self.navhindl()
        Doc.document.add_page_break()
        docx.document.add_heading('NEWS REPORT: SIFY' + var,0)
        self.sifyl()
        Doc.document.add_page_break()
        docx.document.add_heading('NEWS REPORT: NDTV' + var,0)
        self.ndtvl()
        Doc.document.add_page_break()
        docx.document.add_heading('REPORT: IndianExpress' +var, 0)
        self.indian_expressl()
        Doc.document.add_page_break()
        docx.document.add_heading('NEWS REPORT: SCROLL' + var,0)
        self.scll()
        Doc.document.add_page_break()
        docx.document.add_heading('CENTRAL CHRONICAL' + var,0)
        self.ccl()
        Doc.document.add_page_break()
        docx.document.add_heading('SUDERSHAN' + var,0)
        self.ssnl()


        Doc.document.save(str(var1.get()+var2.get()) + ".docx")
        #os.startfile("C:/Users/Shubham/PycharmProjects/News/" + var + ".docx")
        def opendocx():
            os.startfile("G:/Projects Python/News/" + var + ".docx")
        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/"+str(var)+'.docx', 'rb')
            document = Document(f)
            f.close()
        popupmsg(var)"""

#??????????????????MAIN FUNCTION????????????????
    def main(self):
        var = str(var1.get())
        var3 = str(var2.get())
        docx.document.add_heading('DainikBhaskar' + " ", 0)
        self.dbl()
        Doc.document.add_page_break()
        #docx.document.add_heading('NEWS REPORT: TheHindu' + " ", 0)
        #self.thl()
        #Doc.document.add_page_break()
        docx.document.add_heading('NEWS REPORT: HindustanTimes' + " ", 0)
        self.htl()
        Doc.document.add_page_break()
        #docx.document.add_heading('NEWS REPORT: Camera24.in' + " ", 0)

        #Doc.document.add_page_break()
        docx.document.add_heading('NEWS REPORT: AajTak' +  " ", 0)
        self.aajl()
        Doc.document.add_page_break()
        #docx.document.add_heading('NEWS REPORT: firstpost' + " ", 0)
        #self.fpl()
        #Doc.document.add_page_break()
        #docx.document.add_heading('NEWS REPORT: ALT NEWS' + " ", 0)
        #self.altl()
        #Doc.document.add_page_break()
        #docx.document.add_heading('NEWS REPORT: ANI NEWS' + " ", 0)
        #self.anil()
        #Doc.document.add_page_break()
        docx.document.add_heading('NEWS REPORT: NavHind' + " ", 0)
        self.navhindl()
        Doc.document.add_page_break()
        #docx.document.add_heading('NEWS REPORT: SIFY' + " ", 0)
        #self.sifyl()
        #Doc.document.add_page_break()
        docx.document.add_heading('NEWS REPORT: NDTV' + " ", 0)
        self.ndtvl()
        Doc.document.add_page_break()
        docx.document.add_heading('REPORT: IndianExpress' + " ", 0)
        self.indian_expressl()
        #Doc.document.add_page_break()
        #docx.document.add_heading('NEWS REPORT: SCROLL' + " ", 0)
        #self.scll()
        Doc.document.add_page_break()
        docx.document.add_heading('CENTRAL CHRONICAL' + " ", 0)
        self.ccl()
        Doc.document.add_page_break()
        docx.document.add_heading('SUDERSHAN' + " ", 0)
        self.ssnl()
        docx.document.add_heading("Hindi FirstPost"+" ",0)
        self.hfpl()

        #self.fetchlink()


        Doc.document.save(str(var1.get() + var2.get()) + ".docx")

        # os.startfile("C:/Users/Shubham/PycharmProjects/News/" + var + ".docx")
        def opendocx():
            os.startfile("G:/Projects Python/News/" + var+var3+ ".docx")

        def popupmsg(msg):
            popup = Tk()
            popup.wm_title("Report Generated")
            label = Label(popup, text="Your report has been generated would you like to open it ")
            label.pack(side="top", fill="x", pady=10)
            B1 = Button(popup, text="Open", command=opendocx)
            B1.pack()

            B2 = Button(popup, text="Close", command=popup.destroy)
            B2.pack()
            popup.mainloop()

        def openfile():
            f = open("C:/Users/Shubham/PycharmProjects/News/" + str(var)+str(var3) + '.docx', 'rb')
            document = Document(f)
            f.close()

        popupmsg(var)
    #==============THE_HINDU LINK GENERATOR==========
    def thl(self):
        var = str(var1.get())
        var3 = str(var2.get())
        resp = urllib.request.urlopen("https://www.thehindu.com/search/?q="+var+"+"+var3+"&order=DESC&sort=publishdate")
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 50) and (link['href'].find(var)!=-1) ):
                if((link['href'].find('whatsapp') == -1) and (link['href'].find('video')== -1)):
                    print(link['href'])
                    Doc.document.add_paragraph(
                        link['href'], style='List Number'
                    )
                    Doc.document.save(str(var1.get()) + ".docx")
                    #link1 = (link['href'])
                    #self.dba(link1)


#================HINDUSTAN TIMES ARTICLE FETCHER==================
    def hta(self, url):


        resp = requests.get(url)

        if resp.status_code == 200:
            print("Successfully opened the web page")
            print("The news are as follow :-\n")

            soup = BeautifulSoup(resp.text, 'html.parser')

            l = soup.find("div", {"class": "story-details"})
            for i in l.findAll("p"):
                print(i.text)
                Doc.document.add_paragraph(
                    i.text
                )


        else:
            print("Error")



#----------------------------HINDUSTAN_TIMES  LINKS----------------------
    def htl(self):
        var = str(var1.get())
        var3 = str(var2.get())
        resp = urllib.request.urlopen("https://www.bing.com/news/search?q=site%3Ahindustantimes.com++" + var+"+"+var3)
        soup = BeautifulSoup(resp, features="lxml", from_encoding=resp.info().get_param('charset'))
        for link in soup.find_all('a', href=True):
            if ((link['href'].find('http') != -1) and (len(link['href']) > 50) and (link['href'].find(var)!=-1) ):
                if((link['href'].find('whatsapp') == -1) and (link['href'].find('video')== -1)):
                    print(link['href'])
                    p=Doc.document.add_paragraph( style='List Number'

                    )
                    p.add_run(link['href']).bold = True
                    link1 = (link['href'])
                    try:
                        self.hta(link1)
                    except:
                        print("Article fetching is restricted")


                #self.news(link['href'])
                Doc.document.add_page_break()

                Doc.document.save(str(var1.get())+".docx")





    def report_for_links(self):
        document = Document()
        document.add_heading('Links for '+Doc.var1, 0)



var1= str(" ")
docx = Doc(var1)


top = Tk()

top.geometry("530x600")
top.title('SMART News Report Generator')
top.resizable(width=False, height=False)

""" This code is for Background Image for better interaction 

photo = PhotoImage(file = "bg.png")
w = Label(top, image=photo)
w.place(x=0, y=0)
"""

img = PhotoImage(file="Icons\\bg.png")
panel = Label(top, image=img)
panel.place( x=0, y=0)



photo=PhotoImage(file="Icons\\ndtv.png")
b = Button(top, image=photo, command=docx.ndtvll, height=65, width=65, bg="white")
b.place(x = 20, y = 220)

photo1=PhotoImage(file="Icons\\bbc.png")
b = Button(top, image=photo1, command=docx.bbcl, height=65, width=65, bg="red")
b.place(x = 123, y = 220)

#------------------- HINDUSTAN TIMES.COM------------------------
photo2=PhotoImage(file="Icons\\ht.png")
b = Button(top, image=photo2, command=docx.htll, height=65, width=65, bg="red")
b.place(x = 223, y = 220)

photo13=PhotoImage(file="Icons\\fp.png")
b = Button(top, image=photo13, command=docx.fpll, height=65, width=65, bg="white")
b.place(x = 323, y = 220)


photo14=PhotoImage(file="Icons\\alt.png")
b = Button(top, image=photo14, command=docx.altll, height=65, width=65, bg="white")
b.place(x = 430, y = 220)

#---------------------TIMES_NOW.COM---------------------------------
photo3=PhotoImage(file="Icons\\tn.png")
b = Button(top, image=photo3, command=callback, height=65, width=65, bg="white")
b.place(x = 20, y = 320)

#==============Navhind Times
photo15=PhotoImage(file="Icons\\nht.png")
b = Button(top, image=photo15, command=docx.navhindll, height=65, width=65, bg="white")
b.place(x = 123, y = 320)


photo16=PhotoImage(file="Icons\\ss.png")
b = Button(top, image=photo16, command=docx.sifyll, height=65, width=65, bg="white")
b.place(x = 430, y = 320)


photo4=PhotoImage(file="Icons\\ani.png")
b = Button(top, image=photo4, command=docx.anil, height=65, width=65, bg="white")
b.place(x = 223, y = 320)

photo5=PhotoImage(file="Icons\\dna.png")
b = Button(top, image=photo5, command=docx.fetchlink, height=65, width=65, bg="white")
b.place(x = 323, y = 320)

Label(top, text='SEARCH').place(x =20, y = 65)
var1 = StringVar()
var2 = StringVar()
textbox = Entry(top, textvariable=var1, width=10, font=("Helvetica", 16), cursor="dot")
textbox.focus_set()
textbox.place(x=100, y=60)

Label(top, text='+').place(x =235, y = 65)
#Label(top, text='+').place(x =235, y = 65)


textbox1 = Entry(top, textvariable=var2, width=10, font=("Helvetica", 16), cursor="dot")
textbox1.focus_set()
textbox1.place(x=260, y=60)


photo6=PhotoImage(file="Icons\\search.png")
b = Button(top, image=photo6, command=docx.main, height=30, width=30, bg="white")
b.place(x = 480, y = 56)

photo7=PhotoImage(file="Icons\\db.png")
b = Button(top, image=photo7, command=docx.dbll, height=65, width=65, bg="white")
b.place(x = 20, y = 420)


photo8=PhotoImage(file="Icons\\th.png")
b = Button(top, image=photo8, command=docx.thll, height=65, width=65, bg="white")
b.place(x = 123, y = 420)

photo9=PhotoImage(file="Icons\\patrika.png")
b = Button(top, image=photo9, command=docx.pkll, height=65, width=65, bg="white")
b.place(x = 223, y = 420)

photo10=PhotoImage(file="Icons\\c24.png")
b = Button(top, image=photo10, command=docx.c4ll, height=65, width=65, bg="white")
b.place(x = 323, y = 420)

photo11=PhotoImage(file="Icons\\aajtak.png")
b = Button(top, image=photo11, command=docx.aajll, height=65, width=65, bg="white")
b.place(x = 430, y = 420)

photo12=PhotoImage(file="Icons\\cc.png")
b = Button(top, image=photo12, command=docx.ccll, height=65, width=65, bg="white")
b.place(x = 20, y = 520)

photo83=PhotoImage(file="Icons\\ie.png")
b = Button(top, image=photo83, command=docx.indian_expressll, height=65, width=65, bg="white")
b.place(x = 123, y = 520)

photo93=PhotoImage(file="Icons\\scroll.png")
b = Button(top, image=photo93, command=docx.sclll, height=65, width=65, bg="white")
b.place(x = 223, y = 520)

photo103=PhotoImage(file="Icons\\sd.png")
b = Button(top, image=photo103, command=docx.ssnll, height=65, width=65, bg="white")
b.place(x = 323, y = 520)

photo113=PhotoImage(file="Icons\\zee.png")
b = Button(top, image=photo113, command=docx.aajll, height=65, width=65, bg="white")
b.place(x = 430, y = 520)

txt = "News Article Fetcher made by :" \
      "SHUBHAM SINGH RATHORE" \
      "                                                     " \
      "                                           "
T = Text(top, height=5, width=60)
T.place(x=20, y=100)
eer = " "
T.insert(END , eer)


def fo():
    f = open('ram.docx', 'rb')
    document = Document(f)
    f.close()

#b1 = Button(top, command = fo, height=30, width=30 ).place(x = 420, y= 430)



top.mainloop()